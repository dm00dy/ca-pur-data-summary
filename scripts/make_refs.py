"""Generate references.docx and toxicity_sources.csv."""
import csv
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path(__file__).parent

# ── helpers ──────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, runs, space_after=6):
    """Add a paragraph from a list of (text, bold, italic) tuples."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.left_indent = Inches(0)
    for text, bold, italic in runs:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
    return p

def add_hanging(doc, runs, space_after=6):
    """Paragraph with 0.5-inch hanging indent (standard APA reference format)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.left_indent = Inches(0.5)
    for text, bold, italic in runs:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
    return p

def note(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(10)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.italic = True
    return p

# ── document ─────────────────────────────────────────────────────────────────

doc = Document()

# Default style tweaks
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

doc.add_heading("References", level=1)

# ── Section: Toxicity databases ───────────────────────────────────────────────

doc.add_heading("Toxicity Databases and Supporting Literature", level=2)

# 1. PPDB
add_hanging(doc, [
    ("Lewis, K. A., Tzilivakis, J., Warner, D., & Green, A. (2016). "
     "An international database for pesticide risk assessments and management. ", False, False),
    ("Human and Ecological Risk Assessment: An International Journal, 22", False, True),
    ("(4), 1050–1064. ", False, False),
    ("https://doi.org/10.1080/10807039.2015.1133242", False, False),
])
note(doc, "Primary citation for the Pesticide Properties DataBase (PPDB). "
          "Individual chemical records are listed in the accompanying CSV table.")

# 2. BPDB
add_hanging(doc, [
    ("University of Hertfordshire Agriculture and Environment Research Unit (AERU). "
     "(2016). ", False, False),
    ("Biopesticide Properties DataBase (BPDB)", False, True),
    (" [Database]. University of Hertfordshire, Hatfield, UK. "
     "Retrieved May 2026 from https://sitem.herts.ac.uk/aeru/bpdb/", False, False),
])
note(doc, "Used for spinosad (Record 596) — spinosad is not indexed in the PPDB.")

# 3. ECOTOX
add_hanging(doc, [
    ("U.S. Environmental Protection Agency (EPA). (2024). ", False, False),
    ("ECOTOXicology Knowledgebase (ECOTOX)", False, True),
    (" [Database]. National Center for Computational Toxicology, "
     "Research Triangle Park, NC. Retrieved May 2026 from "
     "https://cfpub.epa.gov/ecotox/", False, False),
])
note(doc, "Used for cross-check of beta-cyfluthrin avian acute oral LD50; "
          "supporting reference PMC3090753.")

# 4. APVMA
add_hanging(doc, [
    ("Australian Pesticides and Veterinary Medicines Authority (APVMA). (n.d.). ", False, False),
    ("Public release summary on the evaluation of the new active thiacloprid", False, True),
    (" [Regulatory document]. APVMA, Kingston, ACT. Retrieved May 2026 from "
     "https://www.apvma.gov.au/sites/default/files/publication/14066-prs-thiacloprid.pdf",
     False, False),
])
note(doc, "Used for thiacloprid avian acute oral LD50 (Northern Bobwhite, 2716 mg/kg, "
          "OECD up-and-down method, 95% CI 1,540–248,553 mg/kg). PPDB contains "
          "only a non-guideline canary contact value.")

# 5. Carbaryl literature
add_hanging(doc, [
    ("Norgaard, K. B., & Cedergreen, N. (ca. 2010). "
     "[Full citation to be verified via PubMed PMID 26549316]. ", False, False),
    ("Environmental Toxicology and Chemistry.", False, True),
])
note(doc, "Used for carbaryl Daphnia magna 48-h EC50 = 7.47 μg/L. PPDB contains "
          "only D. pulex 48-h EC50 and D. magna 21-d NOEC for this compound. "
          "VERIFY full title and year at https://pubmed.ncbi.nlm.nih.gov/26549316/ "
          "before submission.")

# ── Section: Pesticide use data ───────────────────────────────────────────────

doc.add_heading("Pesticide Use Data", level=2)

add_hanging(doc, [
    ("California Department of Pesticide Regulation (CDPR). (2024). ", False, False),
    ("Pesticide Use Reporting (PUR)", False, True),
    (" [Database]. CDPR, Sacramento, CA. Retrieved 2015–2023 from "
     "https://www.cdpr.ca.gov/docs/pur/purmain.htm", False, False),
])
note(doc, "Analysis restricted to agricultural sites (site_code < 65,000 per CDPR "
          "site-code conventions); calendar years 2015–2023.")

# ── Footer note ───────────────────────────────────────────────────────────────

doc.add_paragraph()
p = doc.add_paragraph(
    "Note: PPDB record URLs follow the pattern "
    "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/{N}.htm; "
    "BPDB record URLs follow https://sitem.herts.ac.uk/aeru/bpdb/Reports/{N}.htm. "
    "Per-chemical record numbers, test species, and sourcing flags are provided "
    "in the accompanying file toxicity_sources.csv."
)
p.paragraph_format.space_before = Pt(12)
for run in p.runs:
    run.font.size = Pt(10)

out_docx = OUT_DIR / "references.docx"
doc.save(out_docx)
print(f"Wrote {out_docx}")

# ── CSV ───────────────────────────────────────────────────────────────────────

rows = [
    ["chemical", "class", "avian_source", "avian_url", "aquatic_source", "aquatic_url", "flags"],
    ["acephate",          "Organophosphates",   "PPDB #9",    "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/9.htm",    "PPDB #9",    "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/9.htm",    "Mallard surrogate (avian)"],
    ["acetamiprid",       "Other neonics",      "PPDB #11",   "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/11.htm",   "PPDB #11",   "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/11.htm",   "Mallard surrogate (avian)"],
    ["aldicarb",          "Carbamates",         "PPDB #19",   "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/19.htm",   "PPDB #19",   "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/19.htm",   "Mallard surrogate (avian)"],
    ["azinphos-methyl",   "Organophosphates",   "PPDB #51",   "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/51.htm",   "PPDB #51",   "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/51.htm",   ""],
    ["beta-cyfluthrin",   "Pyrethroids",        "ECOTOX (PMC3090753)", "https://cfpub.epa.gov/ecotox/", "PPDB #74",   "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/74.htm",   "PPDB avian value non-guideline (formulated product; canary); ECOTOX Bobwhite >2000 mg/kg used"],
    ["bifenthrin",        "Pyrethroids",        "PPDB #78",   "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/78.htm",   "PPDB #78",   "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/78.htm",   ""],
    ["carbaryl",          "Carbamates",         "PPDB #115",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/115.htm",  "PMID 26549316", "https://pubmed.ncbi.nlm.nih.gov/26549316/",            "Avian ceiling >2000 mg/kg (mallard); PPDB has no D. magna 48h EC50 — literature value used for aquatic"],
    ["carbofuran",        "Carbamates",         "PPDB #118",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/118.htm",  "PPDB #118",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/118.htm",  "Mallard surrogate (avian)"],
    ["chlorantraniliprole","Diamides",           "PPDB #1138", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/1138.htm", "PPDB #1138", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/1138.htm", "Avian ceiling >2250 mg/kg"],
    ["chlorpyrifos",      "Organophosphates",   "PPDB #154",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/154.htm",  "PPDB #154",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/154.htm",  ""],
    ["clothianidin",      "Restricted neonics", "PPDB #171",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/171.htm",  "PPDB #171",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/171.htm",  "Aquatic ceiling >40000 ug/L"],
    ["cyantraniliprole",  "Diamides",           "PPDB #1662", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/1662.htm", "PPDB #1662", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/1662.htm", "Avian ceiling >2250 mg/kg; aquatic ceiling >20 ug/L"],
    ["cyfluthrin",        "Pyrethroids",        "PPDB #192",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/192.htm",  "PPDB #192",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/192.htm",  "Avian ceiling >2000 mg/kg"],
    ["cypermethrin",      "Pyrethroids",        "PPDB #197",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/197.htm",  "PPDB #197",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/197.htm",  "Avian ceiling >9520 mg/kg"],
    ["deltamethrin",      "Pyrethroids",        "PPDB #205",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/205.htm",  "PPDB #205",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/205.htm",  "Avian ceiling >2250 mg/kg"],
    ["diazinon",          "Organophosphates",   "PPDB #212",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/212.htm",  "PPDB #212",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/212.htm",  "Mallard surrogate (avian)"],
    ["dimethoate",        "Organophosphates",   "PPDB #244",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/244.htm",  "PPDB #244",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/244.htm",  ""],
    ["dinotefuran",       "Restricted neonics", "PPDB #1195", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/1195.htm", "PPDB #1195", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/1195.htm", "Avian ceiling >2000 mg/kg (Japanese quail); aquatic ceiling >968300 ug/L"],
    ["esfenvalerate",     "Pyrethroids",        "PPDB #269",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/269.htm",  "PPDB #269",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/269.htm",  ""],
    ["fenpropathrin",     "Pyrethroids",        "PPDB #306",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/306.htm",  "PPDB #306",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/306.htm",  "Mallard surrogate (avian); aquatic ceiling >0.53 ug/L"],
    ["flubendiamide",     "Diamides",           "PPDB #1132", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/1132.htm", "PPDB #1132", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/1132.htm", "Avian ceiling >2000 mg/kg"],
    ["flupyradifurone",   "Next-gen systemics", "PPDB #2620", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/2620.htm", "PPDB #2620", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/2620.htm", "Aquatic ceiling >77600 ug/L"],
    ["gamma-cyhalothrin", "Pyrethroids",        "PPDB #369",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/369.htm",  "PPDB #369",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/369.htm",  "Avian ceiling >2000 mg/kg (mallard)"],
    ["imidacloprid",      "Restricted neonics", "PPDB #397",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/397.htm",  "PPDB #397",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/397.htm",  "Japanese quail surrogate (avian)"],
    ["lambda-cyhalothrin","Pyrethroids",        "PPDB #415",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/415.htm",  "PPDB #415",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/415.htm",  "Avian ceiling >3950 mg/kg (mallard)"],
    ["malathion",         "Organophosphates",   "PPDB #421",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/421.htm",  "PPDB #421",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/421.htm",  ""],
    ["methidathion",      "Organophosphates",   "PPDB #456",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/456.htm",  "PPDB #456",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/456.htm",  "Mallard surrogate (avian)"],
    ["methomyl",          "Carbamates",         "PPDB #458",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/458.htm",  "PPDB #458",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/458.htm",  ""],
    ["naled",             "Organophosphates",   "PPDB #480",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/480.htm",  "PPDB #480",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/480.htm",  "Avian ceiling >52 mg/kg (mallard); low test dose — actual LD50 may be higher"],
    ["nitenpyram",        "Other neonics",      "PPDB #1612", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/1612.htm", "PPDB #1612", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/1612.htm", "Mallard surrogate (avian)"],
    ["oxamyl",            "Carbamates",         "PPDB #498",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/498.htm",  "PPDB #498",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/498.htm",  "Mallard surrogate (avian); highly toxic to birds (3.16 mg/kg)"],
    ["oxydemeton-methyl", "Organophosphates",   "PPDB #501",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/501.htm",  "PPDB #501",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/501.htm",  ""],
    ["permethrin",        "Pyrethroids",        "PPDB #515",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/515.htm",  "PPDB #515",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/515.htm",  "Avian ceiling >9800 mg/kg (mallard)"],
    ["phorate",           "Organophosphates",   "PPDB #519",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/519.htm",  "PPDB #519",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/519.htm",  "Mallard surrogate (avian)"],
    ["phosmet",           "Organophosphates",   "PPDB #521",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/521.htm",  "PPDB #521",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/521.htm",  "Avian ceiling >1068 mg/kg (mallard)"],
    ["spinosad",          "Spinosyns",          "BPDB #596",  "https://sitem.herts.ac.uk/aeru/bpdb/Reports/596.htm",     "BPDB #596",  "https://sitem.herts.ac.uk/aeru/bpdb/Reports/596.htm",     "BPDB not PPDB; avian ceiling >2000 mg/kg (mallard); aquatic ceiling >1000 ug/L"],
    ["spinetoram",        "Spinosyns",          "PPDB #1144", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/1144.htm", "PPDB #1144", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/1144.htm", "Avian ceiling >2250 mg/kg (mallard)"],
    ["sulfoxaflor",       "Next-gen systemics", "PPDB #1669", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/1669.htm", "PPDB #1669", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/1669.htm", "Aquatic ceiling >399000 ug/L"],
    ["tau-fluvalinate",   "Pyrethroids",        "PPDB #608",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/608.htm",  "PPDB #608",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/608.htm",  "Avian ceiling >2510 mg/kg; aquatic 8.9 ug/L is UK CRD regulatory study — 3-source conflict documented in toxicity_lookup.csv notes"],
    ["tetraniliprole",    "Diamides",           "PPDB #3190 (class estimate)", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/3190.htm", "PPDB #3190", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/3190.htm", "FLAG: avian LD50 is class-based estimate — no empirical data in PPDB; verify from EPA registration review before submission"],
    ["thiacloprid",       "Other neonics",      "APVMA PRS",  "https://www.apvma.gov.au/sites/default/files/publication/14066-prs-thiacloprid.pdf", "PPDB #630", "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/630.htm", "No oral LD50 in PPDB; APVMA value used (Bobwhite 2716 mg/kg OECD up-and-down); aquatic ceiling >85100 ug/L"],
    ["thiamethoxam",      "Restricted neonics", "PPDB #631",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/631.htm",  "PPDB #631",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/631.htm",  "Mallard surrogate (avian); aquatic ceiling >100000 ug/L"],
    ["tralomethrin",      "Pyrethroids",        "PPDB #647",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/647.htm",  "PPDB #647",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/647.htm",  "Avian ceiling >2510 mg/kg (Phasianidae family — non-standard surrogate; no Bobwhite or mallard value found); aquatic from L3 reference source"],
    ["zeta-cypermethrin", "Pyrethroids",        "PPDB #682",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/682.htm",  "PPDB #682",  "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/682.htm",  "Avian ceiling >5124 mg/kg (mallard)"],
]

out_csv = OUT_DIR / "toxicity_sources.csv"
with open(out_csv, "w", newline="", encoding="utf-8") as f:
    writer = csv.writer(f)
    writer.writerows(rows)
print(f"Wrote {out_csv}")
