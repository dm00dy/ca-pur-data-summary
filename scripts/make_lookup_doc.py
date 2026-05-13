"""Generate toxicity_lookup_procedure.docx — reviewer verification SOP."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_PATH = Path(__file__).parent / "toxicity_lookup_procedure.docx"

doc = Document()

# ── default style ─────────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)

# ── helpers ───────────────────────────────────────────────────────────────────

def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def body(text, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullets(items, indent=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
        p.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            # (text, bold_prefix)
            prefix, rest = item
            r = p.add_run(prefix)
            r.bold = True
            p.add_run(rest)
        else:
            p.add_run(item)

def numbered(items, indent=0):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
        p.paragraph_format.space_after = Pt(4)
        if isinstance(item, tuple):
            prefix, rest = item
            r = p.add_run(prefix)
            r.bold = True
            p.add_run(rest)
        else:
            p.add_run(item)

def field_box(label, value, note=None):
    """Indented field label + value block, styled like a form entry."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r2 = p.add_run(value)
    r2.font.name = "Courier New"
    r2.font.size = Pt(10)
    if note:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.75)
        p2.paragraph_format.space_after = Pt(6)
        r3 = p2.add_run(note)
        r3.font.size = Pt(9)
        r3.font.italic = True

def warn(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Note: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p.add_run(text)
    for run in p.runs:
        run.font.size = Pt(10)

def spacer():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)

# ══════════════════════════════════════════════════════════════════════════════
# Title
# ══════════════════════════════════════════════════════════════════════════════

h1("Toxicity Value Lookup Procedure")
body(
    "This document describes the step-by-step procedure used to obtain avian "
    "and aquatic invertebrate toxicity values for 44 pesticide active ingredients "
    "included in the California Pesticide Use Reporting (PUR) analysis. "
    "Reviewers may use these instructions to independently verify any individual "
    "value reported in toxicity_sources.csv.",
    space_after=10,
)

body(
    "All values are stored in toxicity_lookup.csv alongside their source URLs. "
    "The procedure below describes how those URLs were queried and which fields "
    "were extracted.",
    space_after=14,
)

# ══════════════════════════════════════════════════════════════════════════════
# Section 1 — Endpoints
# ══════════════════════════════════════════════════════════════════════════════

h2("1. Toxicity Endpoints Used")

body("Two endpoints were extracted per chemical:")

bullets([
    ("Avian acute oral LD50 (mg/kg): ",
     "Single-dose, 14-day mortality. Preferred test species: Northern Bobwhite "
     "(Colinus virginianus). Acceptable surrogate: Mallard duck (Anas platyrhynchos). "
     "Other surrogates are flagged in the notes column of toxicity_lookup.csv."),
    ("Aquatic invertebrate EC50 (μg/L): ",
     "Daphnia magna 48-hour EC50 for immobility (sometimes reported as LC50 "
     "immobility in PPDB). This is the standard Tier-1 regulatory endpoint for "
     "aquatic invertebrate risk assessment."),
])

spacer()
body(
    "These endpoints were chosen to match EPA and EFSA Tier-1 risk assessment "
    "conventions and to maximize cross-chemical comparability.",
    space_after=14,
)

# ══════════════════════════════════════════════════════════════════════════════
# Section 2 — Primary source: PPDB
# ══════════════════════════════════════════════════════════════════════════════

h2("2. Primary Source: Hertfordshire Pesticide Properties DataBase (PPDB)")

body(
    "The PPDB (Lewis et al. 2016) was used as the primary source for all "
    "chemicals except spinosad (see Section 4). Each chemical has a single "
    "curated record at a stable URL.",
    space_after=10,
)

h3("2.1 Locating a chemical record")

numbered([
    'Navigate to the PPDB A–Z index: '
    'https://sitem.herts.ac.uk/aeru/ppdb/en/atoz.htm',
    'Click the first letter of the chemical name. Locate the chemical in the '
    'alphabetical list and click its name.',
    'The record URL will have the form: '
    'https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/{N}.htm  '
    'where N is the PPDB record number. Record this number for citation.',
    'Alternatively, navigate directly if the record number is already known '
    '(e.g., from toxicity_sources.csv).',
])

spacer()
h3("2.2 Extracting the avian LD50")

body("On the chemical record page, locate the section titled \"Ecotoxicological data\" "
     "or \"Ecotoxicology.\" Find the subsection for avian toxicity.")

body("Extract the following field:")
field_box("Field name", "Acute oral LD50",
          "Reported in mg/kg. Prefer the Bobwhite (Colinus virginianus) value. "
          "If only a mallard (Anas platyrhynchos) value is present, use it and "
          "flag the surrogate species in your notes.")

body("If the value is reported as \">N\" (e.g., \">2000\"):")
bullets([
    "Record N as the value (conservative floor).",
    "Set the avian_ceiling flag to TRUE in toxicity_lookup.csv.",
    "Do not interpret \">N\" as meaning the LD50 is N; the true LD50 exceeds N.",
])

warn(
    "PPDB sometimes lists a contact LD50 rather than an oral LD50 for avian toxicity. "
    "Verify that the field explicitly says \"acute oral\" — contact values are not "
    "interchangeable. If only a contact value is available, flag the record and "
    "consult ECOTOX (Section 3) for an oral value."
)

h3("2.3 Extracting the aquatic EC50")

body("In the same Ecotoxicological data section, locate the subsection for "
     "aquatic invertebrates.")

body("Extract the following field:")
field_box("Field name", "Daphnia EC50 (immobility) 48h",
          "PPDB may label this 'LC50 (immobility)' or 'EC50 immobility' — "
          "these refer to the same endpoint. Units in PPDB are mg/L; "
          "multiply by 1000 to convert to μg/L before recording.")

warn(
    "PPDB reports aquatic toxicity in mg/L, not μg/L. A value of 0.00023 mg/L "
    "must be recorded as 0.23 μg/L. Always convert before entering the CSV."
)

body("If the value is reported as \">N\":")
bullets([
    "Record N × 1000 as the value in μg/L (conservative floor).",
    "Set the aquatic_ceiling flag to TRUE.",
])

spacer()

# ══════════════════════════════════════════════════════════════════════════════
# Section 3 — Cross-check: ECOTOX
# ══════════════════════════════════════════════════════════════════════════════

h2("3. Cross-Check Source: EPA ECOTOX Knowledgebase")

body(
    "ECOTOX (https://cfpub.epa.gov/ecotox/) was used as a cross-check when "
    "one of the following conditions was met:",
    space_after=6,
)
bullets([
    "The PPDB avian value uses a non-standard test species and no Bobwhite or "
    "mallard value is available.",
    "The PPDB value is a contact LD50 rather than an oral LD50.",
    "The PPDB value differs from expected class-level values by more than one "
    "order of magnitude.",
    "The PPDB aquatic value cites only Level 3 (reference book) sources rather "
    "than primary study data.",
])

spacer()
h3("3.1 ECOTOX query procedure")

numbered([
    "Navigate to https://cfpub.epa.gov/ecotox/ and click \"Search ECOTOX.\"",
    ("Chemical: ", "Enter the active ingredient name in the \"Chemical\" field. "
     "Use the CAS number if name matching returns no results."),
    ("Species: ", "Filter by species group \"Birds\" (for avian) or "
     "\"Invertebrates\" (for aquatic). Preferred species are listed in Section 1."),
    ("Endpoint: ", "Filter by endpoint \"LD50\" (avian) or \"EC50\" (aquatic). "
     "For aquatic, also filter by \"immobility\" or \"mortality\" effect."),
    ("Duration: ", "Filter by test duration 14d (avian) or 48h (aquatic)."),
    "Review results. Prefer guideline studies (OECD 205 for avian; OECD 202 "
    "for aquatic) over academic studies when both are available.",
    "Record the selected value and the study citation (author, year, journal) "
    "in the notes column of toxicity_lookup.csv.",
])

spacer()
warn(
    "ECOTOX returns multiple values per chemical. Do not average them without "
    "justification. The EPA-selected value used in the most recent registration "
    "review document is the most defensible single value to cite."
)

# ══════════════════════════════════════════════════════════════════════════════
# Section 4 — Special cases
# ══════════════════════════════════════════════════════════════════════════════

h2("4. Special Cases and Exceptions")

h3("4.1 Spinosad — BPDB instead of PPDB")

body(
    "Spinosad is a biopesticide derived from Saccharopolyspora spinosa and is "
    "not indexed in the standard PPDB. Use the AERU Biopesticide Properties "
    "DataBase (BPDB) instead:"
)
field_box("URL", "https://sitem.herts.ac.uk/aeru/bpdb/Reports/596.htm")
body("The BPDB uses the same field names and layout as PPDB. "
     "Apply the same extraction procedure (Sections 2.2–2.3).", space_after=10)

h3("4.2 Thiacloprid — no PPDB oral LD50")

body(
    "The PPDB record for thiacloprid (Record 630) contains only a canary "
    "contact LD50 (35 mg/kg), which is not the required endpoint. The avian "
    "oral LD50 was sourced from the Australian Pesticides and Veterinary "
    "Medicines Authority (APVMA) public release summary:"
)
field_box("URL", "https://www.apvma.gov.au/sites/default/files/publication/14066-prs-thiacloprid.pdf")
body(
    "The value (Northern Bobwhite, 2716 mg/kg, OECD up-and-down method) "
    "appears in the ecotoxicology section of the APVMA review document. "
    "Note the very wide 95% confidence interval (1,540–248,553 mg/kg), "
    "which reflects a low-precision up-and-down design. The value is "
    "interpreted as \"effectively non-toxic to birds\" for risk classification "
    "purposes.", space_after=10
)

h3("4.3 Tetraniliprole — no empirical avian data")

body(
    "The PPDB record for tetraniliprole (Record 3190) did not contain an avian "
    "LD50 at the time of lookup. A class-based estimate of >2000 mg/kg was "
    "used, consistent with all other diamide insecticides in the dataset "
    "(chlorantraniliprole, cyantraniliprole, flubendiamide — all >2000 mg/kg). "
    "This value is flagged in toxicity_lookup.csv and should be replaced with "
    "an empirical value from the EPA registration review before publication."
)
warn(
    "The tetraniliprole avian estimate is the only non-empirical value in the "
    "dataset. Its contribution to the avian toxicity load analysis should be "
    "reported separately or excluded from any figure that claims full empirical "
    "sourcing."
)

h3("4.4 Carbaryl — aquatic from peer-reviewed literature")

body(
    "The PPDB record for carbaryl (Record 115) contains a Daphnia pulex 48-h "
    "EC50 (6.4 μg/L) and a Daphnia magna 21-day NOEC, but not the required "
    "Daphnia magna 48-h EC50. A value of 7.47 μg/L was sourced from "
    "peer-reviewed literature (PubMed PMID 26549316 — verify full citation "
    "before submission). The literature value is consistent with the PPDB "
    "D. pulex value (6.4 μg/L), providing cross-species corroboration."
)

h3("4.5 Beta-cyfluthrin — PPDB avian value non-guideline")

body(
    "The PPDB record for beta-cyfluthrin (Record 74) reports a canary LD50 "
    "derived from a formulated-product study (Dane et al.) rather than a "
    "guideline acute oral study. This value is not appropriate for comparison "
    "with Bobwhite or mallard values from other chemicals. The avian LD50 "
    "was instead sourced from ECOTOX, which confirms a Bobwhite acute oral "
    "LD50 >2000 mg/kg, consistent with the parent compound cyfluthrin "
    "(PPDB Record 192, >2000 mg/kg Bobwhite). The supporting ECOTOX "
    "reference is PMC3090753."
)

spacer()

# ══════════════════════════════════════════════════════════════════════════════
# Section 5 — Ceiling values
# ══════════════════════════════════════════════════════════════════════════════

h2("5. Handling Ceiling Values (\">N\" Reported Values)")

body(
    "A ceiling value occurs when the test dose range did not span the LD50 or "
    "EC50 — that is, no mortality or immobility reached 50% at any tested dose. "
    "Regulators report this as \">N\" where N is the highest dose tested.",
    space_after=8,
)

body("Handling convention used in this analysis:")
bullets([
    ("Conservative floor: ", "Store N as the value. This understates the safety "
     "margin but never overstates toxicity. Toxicity units (lbs ÷ LD50) will "
     "be overestimated for chemicals with ceiling values."),
    ("Boolean flags: ", "The avian_ceiling and aquatic_ceiling columns in "
     "toxicity_lookup.csv record whether each value is a ceiling. Analyses "
     "sensitive to this assumption should treat ceiling chemicals separately."),
    ("Do not double the value: ", "Using 2×N as a substitute is sometimes "
     "recommended but was not applied here, as it introduces an unjustified "
     "assumption about the shape of the dose-response curve above the tested range."),
])

spacer()

# ══════════════════════════════════════════════════════════════════════════════
# Section 6 — Verification checklist
# ══════════════════════════════════════════════════════════════════════════════

h2("6. Reviewer Verification Checklist")

body(
    "To independently verify a value for any chemical in toxicity_sources.csv:",
    space_after=6,
)

numbered([
    "Open toxicity_sources.csv and locate the row for the chemical of interest.",
    ("avian_url / aquatic_url: ", "Navigate to the URL in the relevant column. "
     "For PPDB records, the URL is of the form "
     "https://sitem.herts.ac.uk/aeru/ppdb/en/Reports/{N}.htm."),
    "Locate the endpoint field as described in Section 2 (avian) or Section 2.3 "
    "(aquatic). Confirm the numeric value and units match what is reported in "
    "toxicity_lookup.csv (avian_ld50_mg_per_kg or aquatic_lc50_ug_per_l columns).",
    "Check the avian_test_species and aquatic_test_species columns in "
    "toxicity_lookup.csv against what the source page reports. Flag any "
    "discrepancy.",
    "If a ceiling flag is TRUE, confirm the source reports the value as \">N\" "
    "and that N matches the stored value.",
    "For entries with non-PPDB sources (thiacloprid, carbaryl aquatic, "
    "beta-cyfluthrin avian, spinosad), refer to Sections 4.2–4.5 for "
    "source-specific guidance.",
    "If the PPDB record has been updated since May 2026 and the value has "
    "changed, note the discrepancy. Regulatory database values are occasionally "
    "revised as new studies are accepted.",
])

spacer()

# ══════════════════════════════════════════════════════════════════════════════
# Section 7 — Summary of non-standard sourcing
# ══════════════════════════════════════════════════════════════════════════════

h2("7. Summary of Non-Standard Sourcing Decisions")

body(
    "The table below summarizes the five chemicals where the standard PPDB "
    "procedure could not be followed, and the rationale for the alternative "
    "source used.",
    space_after=8,
)

# Simple table: 3 columns
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr = table.rows[0].cells
for cell, text in zip(hdr, ["Chemical", "Endpoint", "Issue", "Resolution"]):
    cell.text = text
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(10)

non_std = [
    ("spinosad",        "Both",    "Not indexed in PPDB",
     "BPDB Record 596 (same AERU group, same format)"),
    ("thiacloprid",     "Avian",   "PPDB has canary contact LD50 only — not the required endpoint",
     "APVMA public release summary (Bobwhite 2716 mg/kg, OECD up-and-down)"),
    ("tetraniliprole",  "Avian",   "No empirical data in PPDB at time of lookup",
     "Class-based estimate (>2000 mg/kg); flagged; verify before publication"),
    ("carbaryl",        "Aquatic", "PPDB has D. pulex and D. magna 21d NOEC only — missing D. magna 48h EC50",
     "Peer-reviewed literature (PMID 26549316); consistent with D. pulex value"),
    ("beta-cyfluthrin", "Avian",   "PPDB value from non-guideline formulated-product study (canary contact)",
     "ECOTOX (Bobwhite >2000 mg/kg; PMC3090753); consistent with parent cyfluthrin"),
]

for chem, ep, issue, resolution in non_std:
    row = table.add_row().cells
    for cell, text in zip(row, [chem, ep, issue, resolution]):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9)

spacer()
spacer()

body(
    "All other chemicals (39 of 44) were sourced directly from PPDB using the "
    "standard procedure in Section 2, with no cross-checks required.",
    space_after=4,
)

doc.save(OUT_PATH)
print(f"Wrote {OUT_PATH}")
